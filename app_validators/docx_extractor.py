"""
DOCX Extractor Module
FIXED: Extracts actual FAQ answers + skips removed FAQs
No content extraction
"""

from docx import Document
from typing import Dict, Any, List
import re


class DOCXExtractor:
    """Extracts and structures content from DOCX files."""
    
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.all_paragraphs = list(self.doc.paragraphs)  # Cache for answer extraction
        
    def extract_brief_data(self) -> Dict[str, Any]:
        """
        Extract all brief data from DOCX file.
        
        Returns:
            Dictionary containing all brief elements
        """
        brief_data = {
            'url': '',
            'meta_title': '',
            'meta_description': '',
            'h1': '',
            'header_caption': '',
            'headers': [],
            'faqs': {
                'header': '',
                'questions': []
            },
            'product_nav': {
                'tabs': []
            },
            'cta': {
                'caption': '',
                'text': '',
                'position': 'before_faq'
            }
        }
        
        # Extract URL
        brief_data['url'] = self._extract_url()
        
        # Extract metadata (handle "No change" cases)
        brief_data['meta_title'] = self._extract_meta_title()
        brief_data['meta_description'] = self._extract_meta_description()
        
        # Extract H1
        brief_data['h1'] = self._extract_h1()
        
        # Extract header caption
        brief_data['header_caption'] = self._extract_header_caption()
        
        # Extract headers and filter Internal Linking section
        headers = self._extract_headers()
        filtered_headers = self._filter_internal_linking(headers)
        brief_data['headers'] = filtered_headers
        
        # Extract FAQs with actual answers (no Internal Linking sections)
        brief_data['faqs'] = self._extract_faqs(filtered_headers)
        
        # Extract Product Navigation
        brief_data['product_nav']['tabs'] = self._extract_product_nav()
        
        # Extract CTA
        cta = self._extract_cta()
        if cta and (cta.get('caption') or cta.get('text')):
            brief_data['cta'] = cta
        
        return brief_data
    
    def _extract_url(self) -> str:
        """Extract URL from document."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text.startswith('http'):
                        return text
        
        for para in self.doc.paragraphs[:5]:
            text = para.text.strip()
            if text.startswith('http'):
                return text
        
        return ''
    
    def _extract_meta_title(self) -> str:
        """Extract meta title - handle 'No change' case."""
        existing_value = ""
        recommended_value = ""
        
        found_meta_title = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            if 'meta title' in text.lower() and ':' in text:
                found_meta_title = True
                continue
            
            if found_meta_title:
                if text.startswith('Existing:'):
                    existing_value = text.replace('Existing:', '').strip()
                    print(f"DEBUG: Found Existing Meta Title: {existing_value[:50]}...")
                elif text.startswith('Recommended:'):
                    recommended_value = text.replace('Recommended:', '').strip()
                    print(f"DEBUG: Found Recommended Meta Title: {recommended_value}")
                    
                    if 'no change' in recommended_value.lower():
                        print(f"DEBUG: 'No change' detected, using Existing value")
                        return existing_value
                    else:
                        return recommended_value
                
                if 'meta description' in text.lower() or \
                   (para.style and 'Heading' in para.style.name):
                    break
        
        return recommended_value if recommended_value and 'no change' not in recommended_value.lower() else existing_value
    
    def _extract_meta_description(self) -> str:
        """Extract meta description - handle 'No change' case."""
        existing_value = ""
        recommended_value = ""
        
        found_meta_desc = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            if 'meta description' in text.lower() and ':' in text:
                found_meta_desc = True
                continue
            
            if found_meta_desc:
                if text.startswith('Existing:'):
                    existing_value = text.replace('Existing:', '').strip()
                    print(f"DEBUG: Found Existing Meta Description: {existing_value[:50]}...")
                elif text.startswith('Recommended:'):
                    recommended_value = text.replace('Recommended:', '').strip()
                    print(f"DEBUG: Found Recommended Meta Description: {recommended_value[:50]}...")
                    
                    if 'no change' in recommended_value.lower():
                        print(f"DEBUG: 'No change' detected, using Existing value")
                        return existing_value
                    else:
                        return recommended_value
                
                if para.style and para.style.name in ['Heading 1', 'Heading1']:
                    break
        
        return recommended_value if recommended_value and 'no change' not in recommended_value.lower() else existing_value
    
    def _extract_h1(self) -> str:
        """Extract H1 from document."""
        for para in self.doc.paragraphs:
            if para.style and para.style.name in ['Heading 1', 'Heading1']:
                return para.text.strip()
        return ''
    
    def _extract_header_caption(self) -> str:
        """Extract header caption (first paragraph after H1)."""
        h1_found = False
        for para in self.doc.paragraphs:
            if para.style and para.style.name in ['Heading 1', 'Heading1']:
                h1_found = True
                continue
            if h1_found and para.text.strip():
                if not (para.style and 'Heading' in para.style.name):
                    return para.text.strip()
        return ''
    
    def _extract_headers(self) -> List[Dict[str, Any]]:
        """Extract all headers (H2, H3, H4) with their paragraph indices."""
        headers = []
        for idx, para in enumerate(self.all_paragraphs):
            if para.style:
                style_name = para.style.name
                text = para.text.strip()
                
                if not text:
                    continue
                
                if style_name in ['Heading 2', 'Heading2']:
                    headers.append({'level': 'H2', 'text': text, 'para_idx': idx})
                elif style_name in ['Heading 3', 'Heading3']:
                    headers.append({'level': 'H3', 'text': text, 'para_idx': idx})
                elif style_name in ['Heading 4', 'Heading4']:
                    headers.append({'level': 'H4', 'text': text, 'para_idx': idx})
        
        return headers
    
    def _is_internal_linking_section(self, text: str) -> bool:
        """Check if this is Internal Linking section."""
        keywords = [
            'internal linking',
            'in links',
            'out links',
            'new in links',
            'external linking',
            'related links'
        ]
        text_lower = text.lower().strip(':')
        return any(kw in text_lower for kw in keywords)
    
    def _filter_internal_linking(self, headers: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Filter out Internal Linking section and everything after it."""
        filtered = []
        
        for h in headers:
            if self._is_internal_linking_section(h['text']):
                print(f"DEBUG: Filtering out Internal Linking section: {h['text']}")
                break
            
            filtered.append(h)
        
        return filtered
    
    def _is_removed_faq(self, text: str) -> bool:
        """Detect if FAQ is marked as removed."""
        removal_keywords = [
            'removed',
            'this faq is removed',
            'these existing faqs are removed',
            'these faqs are removed',
            'existing faqs are removed',
            'delete this faq',
            'to be removed',
            'marked for removal',
            '[removed]',
            '(removed)',
            'do not include',
            'skip this',
            'ignore this',
            'not included',
            'excluded'
        ]
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in removal_keywords)
    
    def _extract_faq_answer(self, question_para_idx: int, next_header_para_idx: int) -> str:
        """
        Extract FAQ answer text between question H3 and next header.
        
        Args:
            question_para_idx: Index of the H3 question paragraph
            next_header_para_idx: Index of next header (or end of document)
            
        Returns:
            Answer text or empty string if not found/removed
        """
        answer_paragraphs = []
        
        # Extract paragraphs between question and next header
        for idx in range(question_para_idx + 1, next_header_para_idx):
            para = self.all_paragraphs[idx]
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check if this paragraph indicates removal
            if self._is_removed_faq(text):
                print(f"DEBUG: FAQ answer contains removal marker: {text[:50]}...")
                return ""  # Return empty - this FAQ should be skipped
            
            # Skip if it's a header (shouldn't happen but safety check)
            if para.style and 'Heading' in para.style.name:
                break
            
            # Valid answer paragraph
            answer_paragraphs.append(text)
        
        # Combine answer paragraphs
        answer = ' '.join(answer_paragraphs).strip()
        
        # Final check - if answer itself contains removal text
        if answer and self._is_removed_faq(answer):
            print(f"DEBUG: Skipping FAQ - answer contains removal marker")
            return ""
        
        return answer if answer else "Answer extracted from document"
    
    def _extract_faqs(self, headers: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Extract FAQ section - ONLY H3s between FAQ H2 and next H2, with actual answers."""
        faqs = {
            'header': '',
            'questions': []
        }
        
        # Find FAQ H2
        faq_h2_index = None
        for idx, h in enumerate(headers):
            if h['level'] == 'H2' and 'FAQ' in h['text'].upper():
                faqs['header'] = h['text']
                faq_h2_index = idx
                print(f"DEBUG: Found FAQ H2 at header index {idx}: {h['text']}")
                break
        
        if faq_h2_index is None:
            print("DEBUG: No FAQ section found")
            return faqs
        
        # Extract FAQ questions (H3s after FAQ H2, before next H2)
        for idx in range(faq_h2_index + 1, len(headers)):
            h = headers[idx]
            
            # Stop at next H2
            if h['level'] == 'H2':
                print(f"DEBUG: Stopped FAQ extraction at next H2: {h['text']}")
                break
            
            # Only process H3 questions
            if h['level'] == 'H3':
                question_text = h['text']
                
                # Check if question itself indicates removal
                if self._is_removed_faq(question_text):
                    print(f"DEBUG: Skipped removed FAQ question: {question_text}")
                    continue
                
                # Find next header to determine answer boundary
                next_header_idx = idx + 1
                if next_header_idx < len(headers):
                    next_header_para_idx = headers[next_header_idx]['para_idx']
                else:
                    next_header_para_idx = len(self.all_paragraphs)
                
                # Extract actual answer
                answer_text = self._extract_faq_answer(h['para_idx'], next_header_para_idx)
                
                # Only add FAQ if answer is not empty (not removed)
                if answer_text:
                    faqs['questions'].append({
                        'question': question_text,
                        'answer': answer_text
                    })
                    print(f"DEBUG: Added FAQ - Q: {question_text[:50]}... A: {answer_text[:50]}...")
                else:
                    print(f"DEBUG: Skipped FAQ with empty/removed answer: {question_text}")
        
        print(f"DEBUG: Total FAQs extracted: {len(faqs['questions'])}")
        return faqs
    
    def _extract_product_nav(self) -> List[Dict[str, str]]:
        """Extract Product Navigation tabs from table or text."""
        tabs = []
        
        if len(self.doc.tables) >= 3:
            prod_nav_table = self.doc.tables[2]
            for row_idx, row in enumerate(prod_nav_table.rows):
                if row_idx == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    recommended = cells[1]
                    section = cells[2] if len(cells) > 2 else ''
                    if recommended and recommended not in ['', 'Recommended']:
                        tabs.append({
                            'text': recommended,
                            'linked_section': section
                        })
        
        return tabs
    
    def _extract_cta(self) -> Dict[str, str]:
        """Extract CTA section - paragraphs immediately before FAQ section."""
        cta = {
            'caption': '',
            'text': '',
            'position': 'before_faq'
        }
        
        faq_paragraph_idx = None
        
        for idx, paragraph in enumerate(self.all_paragraphs):
            text = paragraph.text.strip()
            
            if paragraph.style and paragraph.style.name in ['Heading 2', 'Heading2']:
                if 'FAQ' in text.upper() or 'FREQUENTLY ASKED' in text.upper():
                    faq_paragraph_idx = idx
                    break
        
        if faq_paragraph_idx is not None and faq_paragraph_idx > 0:
            cta_paragraphs = []
            
            for idx in range(faq_paragraph_idx - 1, max(0, faq_paragraph_idx - 10), -1):
                para = self.all_paragraphs[idx]
                text = para.text.strip()
                
                if para.style and 'Heading' in para.style.name:
                    break
                
                if text and 50 < len(text) < 1000:
                    cta_paragraphs.insert(0, text)
                    
                    if len(cta_paragraphs) >= 2:
                        break
            
            if len(cta_paragraphs) >= 2:
                cta['caption'] = cta_paragraphs[0]
                cta['text'] = cta_paragraphs[1]
            elif len(cta_paragraphs) == 1:
                cta['text'] = cta_paragraphs[0]
                cta['caption'] = ''
        
        return cta
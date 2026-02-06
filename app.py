"""
Fortinet Brief Validator - Streamlit Application
FIXED: Correct field names for ValidationResult
"""

import streamlit as st
import sys
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

sys.path.insert(0, str(Path(__file__).parent))

from app_validators.validator import AEMBriefReviewer, Status, ValidationResult
from app_validators.docx_extractor import DOCXExtractor
from app_validators.url_extractor import URLExtractor
from config import  OPENAI_API_KEY


def init_session_state():
    """Initialize session state variables."""
    if 'validation_results' not in st.session_state:
        st.session_state.validation_results = None
    if 'brief_data' not in st.session_state:
        st.session_state.brief_data = None
    if 'summary_data' not in st.session_state:
        st.session_state.summary_data = None
    if 'unknown_terms_report' not in st.session_state:
        st.session_state.unknown_terms_report = None


def calculate_summary_data(validation_results, brief_data):
    """Calculate summary statistics from validation results."""
    summary_data = {
        'meta_title': {'recognized': True, 'checked': 1, 'passed': 0, 'failed': 0},
        'meta_description': {'recognized': True, 'checked': 1, 'passed': 0, 'failed': 0},
        'h1': {'recognized': True, 'checked': 1, 'passed': 0, 'failed': 0},
        'header_caption': {'recognized': True, 'checked': 1, 'passed': 0, 'failed': 0},
        'h2': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'h3': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'h4': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'faq_h2': {'recognized': True, 'checked': 1, 'passed': 0, 'failed': 0},
        'faq_questions': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'faq_answers': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'product_nav': {'recognized': True, 'checked': 0, 'passed': 0, 'failed': 0},
        'cta': {'recognized': False, 'checked': 0, 'passed': 0, 'failed': 0}
    }
    
    h2_count = len([h for h in brief_data.get('headers', []) if h['level'] == 'H2'])
    h3_count = len([h for h in brief_data.get('headers', []) if h['level'] == 'H3'])
    h4_count = len([h for h in brief_data.get('headers', []) if h['level'] == 'H4'])
    faq_count = len(brief_data.get('faqs', {}).get('questions', []))
    tab_count = len(brief_data.get('product_nav', {}).get('tabs', []))
    
    summary_data['h2']['checked'] = h2_count
    summary_data['h3']['checked'] = h3_count
    summary_data['h4']['checked'] = h4_count
    summary_data['faq_questions']['checked'] = faq_count
    summary_data['faq_answers']['checked'] = faq_count
    summary_data['product_nav']['checked'] = tab_count
    
    cta = brief_data.get('cta', {})
    if cta.get('text'):
        summary_data['cta']['recognized'] = True
        summary_data['cta']['checked'] = 1
    
    for result in validation_results:
        status_key = 'passed' if result['status'] == Status.ACCEPTED.value else 'failed'
        
        if 'Meta Title' in result['use_case']:
            summary_data['meta_title'][status_key] += 1
        elif 'Meta Description' in result['use_case']:
            summary_data['meta_description'][status_key] += 1
        elif result['use_case'] == 'H1':
            summary_data['h1'][status_key] += 1
        elif 'Header Caption' in result['use_case']:
            summary_data['header_caption'][status_key] += 1
        elif result['use_case'] == 'H2':
            summary_data['h2'][status_key] += 1
        elif result['use_case'] == 'H3':
            summary_data['h3'][status_key] += 1
        elif result['use_case'] == 'H4':
            summary_data['h4'][status_key] += 1
        elif 'FAQ H2' in result['use_case'] or 'FAQ Header' in result['use_case']:
            summary_data['faq_h2'][status_key] += 1
        elif 'FAQ Question' in result['use_case']:
            summary_data['faq_questions'][status_key] += 1
        elif 'FAQ Answer' in result['use_case']:
            summary_data['faq_answers'][status_key] += 1
        elif 'Product Nav' in result['use_case'] or 'Tab' in result['use_case']:
            summary_data['product_nav'][status_key] += 1
        elif 'CTA' in result['use_case']:
            summary_data['cta'][status_key] += 1
    
    return summary_data


def render_validation_summary_table(summary_data):
    rows = []
    components = [
        ("Meta Title", "meta_title"),
        ("Meta Description", "meta_description"),
        ("H1", "h1"),
        ("Header Caption", "header_caption"),
        ("H2 Headers", "h2"),
        ("H3 Headers", "h3"),
        ("H4 Headers", "h4"),
        ("FAQ H2 Header Name", "faq_h2"),
        ("FAQ Questions", "faq_questions"),
        ("FAQ Answers", "faq_answers"),
        ("Product Navigation Tabs", "product_nav"),
        ("CTA Section", "cta")
    ]

    for label, key in components:
        data = summary_data[key]
        status = "N/A" if not data["recognized"] else ("PASS" if data["failed"] == 0 else "FAIL")

        rows.append({
            "Component": label,
            "Recognized": "✓" if data["recognized"] else "✗",
            "Checked": data["checked"],
            "Passed": data["passed"],
            "Failed": data["failed"],
            "Status": status
        })

    df = pd.DataFrame(rows)

    def style_status(val):
        if val == "PASS":
            return "color:#198754; font-weight:bold;"
        if val == "FAIL":
            return "color:#dc3545; font-weight:bold;"
        return "color:#6c757d; font-weight:bold;"

    st.subheader("VALIDATION SUMMARY")
    st.dataframe(df.style.map(style_status, subset=["Status"]), use_container_width=True)


def generate_docx_report(summary_data, validation_results):
    """Generate DOCX report with all validation results."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Fortinet Brief Validation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Timestamp
    timestamp = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section 1: Validation Summary
    doc.add_heading('1. VALIDATION SUMMARY', 1)
    
    # Create summary table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Component', 'Recognized', 'Checked', 'Passed', 'Failed', 'Status']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    components = [
        ("Meta Title", "meta_title"),
        ("Meta Description", "meta_description"),
        ("H1", "h1"),
        ("Header Caption", "header_caption"),
        ("H2 Headers", "h2"),
        ("H3 Headers", "h3"),
        ("H4 Headers", "h4"),
        ("FAQ H2 Header", "faq_h2"),
        ("FAQ Questions", "faq_questions"),
        ("FAQ Answers", "faq_answers"),
        ("Product Nav Tabs", "product_nav"),
        ("CTA Section", "cta")
    ]
    
    for label, key in components:
        data = summary_data[key]
        status = "N/A" if not data["recognized"] else ("PASS" if data["failed"] == 0 else "FAIL")
        
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = "✓" if data["recognized"] else "✗"
        row_cells[2].text = str(data["checked"])
        row_cells[3].text = str(data["passed"])
        row_cells[4].text = str(data["failed"])
        row_cells[5].text = status
    
    doc.add_paragraph()
    
    # Section 2: Failed Items
    failed_items = [r for r in validation_results if r['status'] != Status.ACCEPTED.value]
    
    if failed_items:
        doc.add_heading('2. FAILED ITEMS', 1)
        doc.add_paragraph(f'Total: {len(failed_items)} issues found')
        doc.add_paragraph()
        
        for idx, item in enumerate(failed_items, 1):
            doc.add_heading(f"{idx}. {item['use_case']} - {item['criterion']}", 2)
            doc.add_paragraph(f"Category: {item['category']}")
            doc.add_paragraph(f"Current: {item['location']}")
            doc.add_paragraph(f"Recommended: {item['details']}")
            
            # Add AI info if available (FIXED FIELD NAMES)
            if item.get('ai_result'):
                doc.add_paragraph(f"AI Correction: {item['ai_result']}")
            if item.get('unknown_terms'):
                doc.add_paragraph(f"Unknown Terms: {', '.join(item['unknown_terms'])}")
            
            doc.add_paragraph()
    
    # Section 3: Passed Items
    passed_items = [r for r in validation_results if r['status'] == Status.ACCEPTED.value]
    
    if passed_items:
        doc.add_heading('3. PASSED ITEMS', 1)
        doc.add_paragraph(f'Total: {len(passed_items)} items passed')
        doc.add_paragraph()
        
        for item in passed_items[:20]:
            doc.add_paragraph(f"✓ {item['use_case']}: {item['location'][:80]}...", style='List Bullet')
        
        if len(passed_items) > 20:
            doc.add_paragraph(f"... and {len(passed_items) - 20} more items")
        
        doc.add_paragraph()
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()


def validate_brief(brief_data, openai_api_key=None):
    """Run validation on brief data with optional OpenAI integration."""
    reviewer = AEMBriefReviewer(openai_api_key=openai_api_key)
    results = reviewer.review_brief(brief_data)
    
    results_dict = []
    for r in results:
        result = {
            'use_case': r.use_case,
            'criterion': r.criterion,
            'location': r.location,
            'status': r.status.value,
            'details': r.details,
            'category': r.category
        }
        
        # Add AI fields if present (FIXED FIELD NAMES)
        if r.ai_validated:
            result['ai_validated'] = r.ai_validated
        if r.ai_result:
            result['ai_result'] = r.ai_result
        if r.unknown_terms:
            result['unknown_terms'] = r.unknown_terms
        
        results_dict.append(result)
    
    return results_dict


def collect_unknown_terms(validation_results):
    """Collect all unknown terms from validation results."""
    unknown_by_section = {}
    
    for result in validation_results:
        if result.get('unknown_terms'):
            section = result['use_case']
            unknown_by_section[section] = result['unknown_terms']
    
    return unknown_by_section


def main():
    """Main Streamlit application."""
    
    st.set_page_config(
        page_title="Fortinet Content Brief Validator",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    init_session_state()
    
    # Check API keys
   
    
    # # Show OpenAI status
    # if OPENAI_API_KEY:
    #     st.success("✅ OpenAI AI Validation Enabled")
    # else:
    #     st.warning("⚠️ OpenAI API key not found-Using rule-based validation only")
    
    st.title("📋 Fortinet Content Brief Validator")
    st.header("Upload Document or Enter URL")
    
    if st.session_state.brief_data or st.session_state.validation_results:
        if st.button("🔄 Start New Validation"):
            st.session_state.brief_data = None
            st.session_state.validation_results = None
            st.session_state.summary_data = None
            st.session_state.unknown_terms_report = None
            st.rerun()
    
    input_method = st.radio(
        "Select input method:",
        ["Upload DOCX File", "Enter Live URL"],
        horizontal=True
    )
    
    brief_data = None
    
    if input_method == "Upload DOCX File":
        uploaded_file = st.file_uploader(
            "Upload Fortinet Brief (DOCX)",
            type=['docx'],
            help="Upload a .docx file containing the Fortinet brief"
        )
        
        if uploaded_file is not None:
            temp_path = Path(f"temp_{uploaded_file.name}")
            with open(temp_path, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            
            try:
                with st.spinner("Extracting data from DOCX..."):
                    extractor = DOCXExtractor(str(temp_path))
                    brief_data = extractor.extract_brief_data()
                
                st.success("✅ DOCX file processed successfully")
                temp_path.unlink()
                
            except Exception as e:
                st.error(f"Error processing DOCX: {str(e)}")
                if temp_path.exists():
                    temp_path.unlink()
    
    else:
        url = st.text_input(
            "Enter Fortinet Brief URL",
            placeholder="https://www.fortinet.com/resources/cyberglossary/application-security"
        )
        
        if st.button("Fetch & Extract") and url:
            try:
                with st.spinner("Fetching data from URL..."):
                    extractor = URLExtractor(url)
                    brief_data = extractor.extract_brief_data()
                
                st.success("✅ URL processed successfully")
                
            except Exception as e:
                st.error(f"Error processing URL: {str(e)}")
    
    if brief_data:
        st.session_state.brief_data = brief_data
    
    if st.session_state.brief_data:
        if st.button("🔍 Run Validation", type="primary"):
            with st.spinner("Running validation with AI..."):
                # Pass OpenAI key to validator
                validation_results = validate_brief(
                    st.session_state.brief_data,
                    openai_api_key=OPENAI_API_KEY if OPENAI_API_KEY else None
                )
                
                summary_data = calculate_summary_data(validation_results, st.session_state.brief_data)
                
                # Collect unknown terms
                # unknown_terms_report = collect_unknown_terms(validation_results)
                
                st.session_state.validation_results = validation_results
                st.session_state.summary_data = summary_data
                # st.session_state.unknown_terms_report = unknown_terms_report
            
            st.success("✅ Validation complete!")
    
    if st.session_state.validation_results:
        st.markdown("---")
        st.header("📊 Validation Results")
        
        results = st.session_state.validation_results
        
        # Calculate stats
        total = len(results)
        passed = sum(1 for r in results if r['status'] == Status.ACCEPTED.value)
        failed = total - passed
        pass_rate = (passed / total * 100) if total > 0 else 0
        
        # Check for AI validated items
        ai_items = [r for r in results if r.get('ai_validated')]
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Checks", total)
        with col2:
            st.metric("Passed", passed)
        with col3:
            st.metric("Failed", failed)
        with col4:
            st.metric("Pass Rate", f"{pass_rate:.1f}%")
        
       
        
        if failed == 0:
            st.success("🎉 All validations passed!")
        else:
            st.warning(f"⚠️ {failed} corrections required")
        
     
        st.markdown("---")
        render_validation_summary_table(st.session_state.summary_data)
        
        # Failed Items Table
        if failed > 0:
            st.markdown("---")
            st.subheader("❌ Failed Items")
            
            reviewer = AEMBriefReviewer(openai_api_key=OPENAI_API_KEY)

            
            failed_results = [r for r in results if r['status'] != Status.ACCEPTED.value]
            
            result_objects = []
            for r in failed_results:
                obj = ValidationResult(
                    use_case=r['use_case'],
                    criterion=r['criterion'],
                    location=r['location'],
                    status=Status(r['status']),
                    details=r['details'],
                    category=r['category']
                    # ai_validated=r.get('ai_validated', False),
                    # ai_result=r.get('ai_result', ''),
                    # unknown_terms=r.get('unknown_terms')
                )
                result_objects.append(obj)
            
            failed_table = reviewer.generate_failed_items_table(result_objects)
            
            if not failed_table.empty:
                # Show AI columns if present
                column_config = {
                    "Category": st.column_config.TextColumn("Category", width="small"),
                    "Type": st.column_config.TextColumn("Type", width="medium"),
                    "Current": st.column_config.TextColumn("Current", width="large"),
                    "Fix": st.column_config.TextColumn("Fix", width="large"),
                    "Recommended": st.column_config.TextColumn("Recommended", width="large"),
                }
      
                
                st.data_editor(
                    failed_table,
                    height=600,
                    use_container_width=True,
                    hide_index=True,
                    disabled=True,
                    column_config=column_config
                )
        
        # Passed Items
        passed_items = [r for r in results if r['status'] == Status.ACCEPTED.value]
        if passed_items:
            with st.expander(f"✅ Passed Items ({len(passed_items)})"):
                for item in passed_items:
                    st.write(f"- {item['use_case']}: {item['location'][:80]}...")
        
        # Download Report Button
        st.markdown("---")
        st.subheader("📥 Download Validation Report")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Generate DOCX report
            docx_data = generate_docx_report(
                st.session_state.summary_data,
                st.session_state.validation_results
            )
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            st.download_button(
                label="📄 Download as DOCX",
                data=docx_data,
                file_name=f"validation_report_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )


if __name__ == "__main__":
    main()